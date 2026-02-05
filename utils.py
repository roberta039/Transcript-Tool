import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import uuid

def generate_session_id() -> str:
    """Generează un ID unic pentru sesiune"""
    return str(uuid.uuid4())[:8]

def get_session_id_from_url() -> str:
    """Obține ID-ul sesiunii din URL sau generează unul nou"""
    query_params = st.query_params
    
    if "session" in query_params:
        return query_params["session"]
    
    return None

def set_session_id_in_url(session_id: str):
    """Setează ID-ul sesiunii în URL"""
    st.query_params["session"] = session_id

def create_word_document(transcription: str, video_name: str, 
                         source_lang: str, target_lang: str) -> BytesIO:
    """Creează un document Word cu transcrierea"""
    
    doc = Document()
    
    # Titlu
    title = doc.add_heading('Transcriere Video', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informații despre fișier
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('📹 Fișier video: ').bold = True
    info_para.add_run(video_name)
    
    info_para2 = doc.add_paragraph()
    info_para2.add_run('🌐 Limba sursă: ').bold = True
    info_para2.add_run(source_lang)
    
    info_para3 = doc.add_paragraph()
    info_para3.add_run('🎯 Limba țintă: ').bold = True
    info_para3.add_run(target_lang)
    
    info_para4 = doc.add_paragraph()
    info_para4.add_run('📅 Data generării: ').bold = True
    info_para4.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))
    
    # Separator
    doc.add_paragraph('─' * 50)
    
    # Transcriere
    doc.add_heading('Transcriere', level=1)
    
    # Împarte transcrierea în paragrafe
    paragraphs = transcription.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text)
            para.paragraph_format.space_after = Pt(6)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    footer = doc.add_paragraph()
    footer.add_run('Generat cu AI Video Transcriber powered by Gemini').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvează în BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def format_timestamp(timestamp_str: str) -> str:
    """Formatează un timestamp pentru afișare"""
    try:
        dt = datetime.fromisoformat(timestamp_str.replace('Z', '+00:00'))
        return dt.strftime("%d.%m.%Y %H:%M")
    except:
        return timestamp_str

def get_status_emoji(status: str) -> str:
    """Returnează emoji pentru status"""
    status_emojis = {
        "active": "✅",
        "expired": "❌",
        "error": "⚠️",
        "completed": "✅",
        "processing": "⏳",
        "failed": "❌"
    }
    return status_emojis.get(status.lower(), "❓")
