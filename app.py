import streamlit as st
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
import tempfile
import os
import time
from io import BytesIO
import json
import re
import requests

# Import Gemini
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    st.error("❌ google-generativeai nu este instalat")
    GEMINI_AVAILABLE = False

# Import python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import pentru YouTube și procesare video
try:
    import yt_dlp
    YTDLP_AVAILABLE = True
except ImportError:
    YTDLP_AVAILABLE = False

try:
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    import io
    GDRIVE_AVAILABLE = True
except ImportError:
    GDRIVE_AVAILABLE = False

# ==================== CONFIGURARE ====================

st.set_page_config(
    page_title="🎬 AI Video Transcriber",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# LIMITE FIȘIERE
MAX_FILE_SIZE_MB = 1000  # Limită maximă aplicație (1GB)
GEMINI_DIRECT_UPLOAD_LIMIT_MB = 200  # Limită pentru upload direct Gemini
YOUTUBE_MAX_DURATION_MINUTES = 120  # Limită durată YouTube (2 ore)

# CSS
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    .main-header h1 {
        margin: 0;
        font-size: 2.5rem;
    }
    .main-header p {
        margin: 0.5rem 0 0 0;
        opacity: 0.95;
    }
    .session-info {
        background-color: #e7f3ff;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #0066cc;
        margin-bottom: 1rem;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffc107;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #dc3545;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #28a745;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    .transcription-box {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #dee2e6;
        max-height: 500px;
        overflow-y: auto;
        font-family: 'Courier New', monospace;
        white-space: pre-wrap;
    }
    .url-input-container {
        background: linear-gradient(135deg, #667eea15 0%, #764ba215 100%);
        padding: 1.5rem;
        border-radius: 12px;
        border: 2px solid #667eea30;
        margin: 1rem 0;
    }
    .feature-card {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 0.5rem 0;
        transition: transform 0.2s;
    }
    .feature-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0,0,0,0.15);
    }
</style>
""", unsafe_allow_html=True)

# ==================== DATABASE ====================

DB_PATH = Path("data")
DB_PATH.mkdir(exist_ok=True)
DB_FILE = DB_PATH / "sessions.db"

def check_and_migrate_database():
    """Verifică și migrează baza de date dacă e nevoie"""
    conn = sqlite3.connect(str(DB_FILE))
    cursor = conn.cursor()
    
    try:
        cursor.execute("PRAGMA table_info(transcriptions)")
        columns = [column[1] for column in cursor.fetchall()]
        
        columns_to_add = {
            'source_language': "TEXT DEFAULT 'Auto-detect'",
            'target_language': "TEXT DEFAULT 'Română'",
            'status': "TEXT DEFAULT 'completed'",
            'file_size_mb': "REAL DEFAULT 0",
            'process_method': "TEXT DEFAULT 'direct'",
            'source_url': "TEXT",
            'source_type': "TEXT DEFAULT 'upload'"
        }
        
        for col_name, col_def in columns_to_add.items():
            if col_name not in columns:
                cursor.execute(f"""
                    ALTER TABLE transcriptions 
                    ADD COLUMN {col_name} {col_def}
                """)
                conn.commit()
            
    except Exception as e:
        pass
    finally:
        conn.close()

def init_database():
    """Inițializează baza de date SQLite"""
    conn = sqlite3.connect(str(DB_FILE))
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sessions (
            session_id TEXT PRIMARY KEY,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS transcriptions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT,
            video_name TEXT,
            source_language TEXT DEFAULT 'Auto-detect',
            target_language TEXT DEFAULT 'Română',
            transcription TEXT,
            status TEXT DEFAULT 'completed',
            file_size_mb REAL DEFAULT 0,
            process_method TEXT DEFAULT 'direct',
            source_url TEXT,
            source_type TEXT DEFAULT 'upload',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (session_id) REFERENCES sessions(session_id)
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT,
            role TEXT,
            content TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (session_id) REFERENCES sessions(session_id)
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS api_key_usage (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            key_index INTEGER,
            status TEXT,
            error_message TEXT,
            used_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    conn.commit()
    conn.close()
    
    check_and_migrate_database()

# ==================== SESSION MANAGEMENT ====================

def generate_session_id():
    return str(uuid.uuid4())[:8]

def get_session_id_from_url():
    params = st.query_params
    return params.get("session", None)

def set_session_id_in_url(session_id):
    st.query_params["session"] = session_id

def init_session():
    url_session_id = get_session_id_from_url()
    
    if url_session_id and session_exists(url_session_id):
        st.session_state.session_id = url_session_id
        
        if 'session_loaded' not in st.session_state:
            st.session_state.messages = get_messages(url_session_id)
            st.session_state.transcriptions = get_transcriptions(url_session_id)
            st.session_state.session_loaded = True
    else:
        if 'session_id' not in st.session_state:
            new_session_id = generate_session_id()
            st.session_state.session_id = new_session_id
            create_session(new_session_id)
            set_session_id_in_url(new_session_id)
            st.session_state.messages = []
            st.session_state.transcriptions = []
            st.session_state.session_loaded = True

def session_exists(session_id):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        cursor.execute("SELECT 1 FROM sessions WHERE session_id = ?", (session_id,))
        exists = cursor.fetchone() is not None
        conn.close()
        return exists
    except:
        return False

def create_session(session_id):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        cursor.execute(
            "INSERT OR IGNORE INTO sessions (session_id) VALUES (?)",
            (session_id,)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Eroare la crearea sesiunii: {e}")

def delete_session_data(session_id):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        cursor.execute("DELETE FROM messages WHERE session_id = ?", (session_id,))
        cursor.execute("DELETE FROM transcriptions WHERE session_id = ?", (session_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Eroare la ștergerea datelor: {e}")

# ==================== API KEY MANAGEMENT ====================

def get_api_keys_from_secrets():
    keys = []
    
    try:
        if "GEMINI_API_KEYS" in st.secrets:
            secret_keys = st.secrets["GEMINI_API_KEYS"]
            if isinstance(secret_keys, list):
                keys.extend(secret_keys)
            elif isinstance(secret_keys, str):
                keys.extend([k.strip() for k in secret_keys.split(",") if k.strip()])
        
        if "GEMINI_API_KEY" in st.secrets:
            single_key = st.secrets["GEMINI_API_KEY"]
            if single_key and single_key not in keys:
                keys.append(single_key)
                
    except Exception as e:
        pass
    
    return keys

def test_api_key(api_key):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content('Say "OK"')
        return True, "✅ Cheie validă"
    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower() or "billing" in error_msg.lower():
            return False, "❌ Cheie expirată (quota/billing)"
        elif "api key" in error_msg.lower():
            return False, "❌ Cheie invalidă"
        else:
            return False, f"❌ Eroare: {error_msg[:100]}"

def get_working_api_key(keys):
    if not keys:
        return None, None, "Nu există chei API configurate"
    
    for i, key in enumerate(keys):
        valid, msg = test_api_key(key)
        if valid:
            return key, i, msg
    
    return None, None, "Toate cheile API sunt expirate sau invalide"

# ==================== DATABASE OPERATIONS ====================

def save_message(session_id, role, content):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO messages (session_id, role, content) VALUES (?, ?, ?)",
            (session_id, role, content)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"Eroare salvare mesaj: {e}")

def get_messages(session_id):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        cursor.execute(
            "SELECT role, content, created_at FROM messages WHERE session_id = ? ORDER BY created_at",
            (session_id,)
        )
        messages = [{"role": row[0], "content": row[1], "time": row[2]} for row in cursor.fetchall()]
        conn.close()
        return messages
    except:
        return []

def save_transcription(session_id, video_name, source_lang, target_lang, transcription, 
                       file_size_mb=0, process_method="direct", source_url="", source_type="upload"):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO transcriptions 
            (session_id, video_name, source_language, target_language, transcription, 
             status, file_size_mb, process_method, source_url, source_type) 
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (session_id, video_name, source_lang, target_lang, transcription, 
              'completed', file_size_mb, process_method, source_url, source_type))
        
        conn.commit()
        transcription_id = cursor.lastrowid
        conn.close()
        return transcription_id
    except Exception as e:
        st.error(f"Eroare salvare transcriere: {e}")
        return None

def get_transcriptions(session_id):
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, video_name, source_language, target_language, 
                   transcription, status, file_size_mb, process_method,
                   source_url, source_type, created_at 
            FROM transcriptions 
            WHERE session_id = ? 
            ORDER BY created_at DESC
        ''', (session_id,))
        
        transcriptions = []
        for row in cursor.fetchall():
            transcriptions.append({
                "id": row[0],
                "video_name": row[1],
                "source_language": row[2],
                "target_language": row[3],
                "transcription": row[4],
                "status": row[5],
                "file_size_mb": row[6],
                "process_method": row[7],
                "source_url": row[8],
                "source_type": row[9],
                "created_at": row[10]
            })
        
        conn.close()
        return transcriptions
    except Exception as e:
        st.error(f"Eroare citire transcrieri: {e}")
        return []

# ==================== URL PROCESSING ====================

def extract_video_id_youtube(url):
    """Extrage ID-ul video din URL YouTube"""
    patterns = [
        r'(?:youtube\.com\/watch\?v=)([\w-]+)',
        r'(?:youtu\.be\/)([\w-]+)',
        r'(?:youtube\.com\/embed\/)([\w-]+)',
        r'(?:youtube\.com\/v\/)([\w-]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def extract_file_id_gdrive(url):
    """Extrage ID-ul fișierului din URL Google Drive"""
    patterns = [
        r'(?:drive\.google\.com\/file\/d\/)([\w-]+)',
        r'(?:drive\.google\.com\/open\?id=)([\w-]+)',
        r'(?:docs\.google\.com\/.*\/d\/)([\w-]+)',
        r'(?:drive\.google\.com\/uc\?id=)([\w-]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def detect_url_type(url):
    """Detectează tipul URL-ului"""
    if not url:
        return None
    
    url_lower = url.lower()
    
    if 'youtube.com' in url_lower or 'youtu.be' in url_lower:
        return 'youtube'
    elif 'drive.google.com' in url_lower or 'docs.google.com' in url_lower:
        return 'gdrive'
    elif url_lower.startswith('http://') or url_lower.startswith('https://'):
        # Verifică dacă e link direct către video
        video_extensions = ['.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv']
        for ext in video_extensions:
            if ext in url_lower:
                return 'direct'
        return 'other'
    
    return None

def get_youtube_info(video_id):
    """Obține informații despre video YouTube"""
    if not YTDLP_AVAILABLE:
        return None, "yt-dlp nu este instalat"
    
    try:
        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
            'extract_flat': True
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(f"https://www.youtube.com/watch?v={video_id}", download=False)
            
            return {
                'title': info.get('title', 'Unknown'),
                'duration': info.get('duration', 0),
                'uploader': info.get('uploader', 'Unknown'),
                'view_count': info.get('view_count', 0),
                'description': info.get('description', '')[:500]
            }, None
    except Exception as e:
        return None, str(e)

def download_youtube_video(video_id, max_size_mb=200, progress_callback=None):
    """Descarcă video YouTube"""
    if not YTDLP_AVAILABLE:
        return None, None, "yt-dlp nu este instalat"
    
    try:
        output_path = tempfile.mktemp(suffix='.mp4')
        
        if progress_callback:
            progress_callback(0.1, "📥 Descărcare video YouTube...")
        
        ydl_opts = {
            'format': f'best[filesize<{max_size_mb}M]/bestvideo[filesize<{max_size_mb}M]+bestaudio/best',
            'outtmpl': output_path,
            'quiet': True,
            'no_warnings': True,
            'extract_flat': False,
            'postprocessors': [{
                'key': 'FFmpegVideoConvertor',
                'preferedformat': 'mp4'
            }]
        }
        
        # Încearcă descărcarea cu limite de calitate descrescătoare
        quality_formats = [
            f'best[height<=720][filesize<{max_size_mb}M]',
            f'best[height<=480][filesize<{max_size_mb}M]', 
            f'worst[filesize<{max_size_mb}M]'
        ]
        
        for format_spec in quality_formats:
            try:
                ydl_opts['format'] = format_spec
                
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    info = ydl.extract_info(f"https://www.youtube.com/watch?v={video_id}", download=True)
                    
                    if os.path.exists(output_path):
                        file_size_mb = os.path.getsize(output_path) / (1024 * 1024)
                        
                        if progress_callback:
                            progress_callback(0.5, f"✅ Video descărcat ({file_size_mb:.1f}MB)")
                        
                        return output_path, info.get('title', 'YouTube Video'), None
                    
            except Exception as e:
                continue
        
        # Dacă nu merge cu video, încearcă doar audio
        if progress_callback:
            progress_callback(0.3, "🎵 Încerc doar audio...")
        
        audio_path = tempfile.mktemp(suffix='.m4a')
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': audio_path,
            'quiet': True,
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'm4a',
                'preferredquality': '128'
            }]
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(f"https://www.youtube.com/watch?v={video_id}", download=True)
            
            if os.path.exists(audio_path):
                if progress_callback:
                    progress_callback(0.7, "✅ Audio descărcat")
                
                return audio_path, info.get('title', 'YouTube Audio'), 'audio_only'
        
        return None, None, "Nu s-a putut descărca video/audio"
        
    except Exception as e:
        return None, None, f"Eroare descărcare: {str(e)}"

def download_gdrive_video(file_id, progress_callback=None):
    """Descarcă video de pe Google Drive"""
    try:
        if progress_callback:
            progress_callback(0.1, "📥 Descărcare de pe Google Drive...")
        
        # URL pentru descărcare directă
        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
        
        # Încearcă descărcarea
        response = requests.get(download_url, stream=True)
        
        # Verifică pentru confirmare (fișiere mari)
        if 'download_warning' in response.content.decode():
            # Extrage token pentru confirmare
            token = None
            for line in response.content.decode().split('\n'):
                if 'download_warning' in line:
                    match = re.search(r'download_warning.*?=([^&]+)', line)
                    if match:
                        token = match.group(1)
                        break
            
            if token:
                download_url = f"https://drive.google.com/uc?export=download&confirm={token}&id={file_id}"
                response = requests.get(download_url, stream=True)
        
        if response.status_code == 200:
            # Salvează fișierul
            output_path = tempfile.mktemp(suffix='.mp4')
            
            total_size = int(response.headers.get('content-length', 0))
            downloaded = 0
            
            with open(output_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        
                        if progress_callback and total_size > 0:
                            progress = 0.1 + (0.8 * downloaded / total_size)
                            progress_callback(progress, f"📥 Descărcat {downloaded/(1024*1024):.1f}MB")
            
            file_size_mb = os.path.getsize(output_path) / (1024 * 1024)
            
            if progress_callback:
                progress_callback(0.9, f"✅ Descărcat ({file_size_mb:.1f}MB)")
            
            return output_path, f"GDrive_{file_id[:8]}.mp4", None
        else:
            return None, None, f"Eroare HTTP: {response.status_code}"
            
    except Exception as e:
        return None, None, f"Eroare descărcare GDrive: {str(e)}"

def download_direct_video(url, progress_callback=None):
    """Descarcă video de la URL direct"""
    try:
        if progress_callback:
            progress_callback(0.1, "📥 Descărcare video...")
        
        response = requests.get(url, stream=True)
        
        if response.status_code == 200:
            # Determină extensia
            content_type = response.headers.get('content-type', '')
            ext = '.mp4'
            if 'video/' in content_type:
                ext = '.' + content_type.split('/')[-1].split(';')[0]
            
            output_path = tempfile.mktemp(suffix=ext)
            
            total_size = int(response.headers.get('content-length', 0))
            downloaded = 0
            
            with open(output_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        
                        if progress_callback and total_size > 0:
                            progress = 0.1 + (0.8 * downloaded / total_size)
                            progress_callback(progress, f"📥 {downloaded/(1024*1024):.1f}/{total_size/(1024*1024):.1f}MB")
            
            file_name = url.split('/')[-1].split('?')[0] or 'direct_video.mp4'
            
            if progress_callback:
                progress_callback(0.9, "✅ Video descărcat")
            
            return output_path, file_name, None
        else:
            return None, None, f"Eroare HTTP: {response.status_code}"
            
    except Exception as e:
        return None, None, f"Eroare descărcare: {str(e)}"

# ==================== VIDEO PROCESSING ====================

SUPPORTED_FORMATS = ['mp4', 'mpeg', 'mov', 'avi', 'mkv', 'webm', 'flv', 'wmv', '3gp', 'ogv']

LANGUAGES = {
    "Română": "Romanian",
    "Engleză": "English",
    "Spaniolă": "Spanish",
    "Franceză": "French",
    "Germană": "German",
    "Italiană": "Italian",
    "Portugheză": "Portuguese",
    "Rusă": "Russian",
    "Chineză": "Chinese",
    "Japoneză": "Japanese",
    "Coreeană": "Korean",
    "Arabă": "Arabic",
    "Hindusă": "Hindi",
    "Turcă": "Turkish",
    "Auto-detect": "auto"
}

def process_and_transcribe(file_path, source_lang, target_lang, api_key, 
                           file_size_mb=0, progress_callback=None, is_audio_only=False):
    """Procesează și transcrie fișierul video/audio"""
    try:
        genai.configure(api_key=api_key)
        
        if is_audio_only:
            # Pentru fișiere audio
            if progress_callback:
                progress_callback(0.3, "🎵 Procesare fișier audio...")
            
            audio_file = genai.upload_file(path=file_path)
            
            if progress_callback:
                progress_callback(0.5, "⏳ Așteptare procesare...")
            
            # Așteaptă procesarea
            max_wait = 60
            wait_time = 0
            while audio_file.state.name == "PROCESSING" and wait_time < max_wait:
                time.sleep(2)
                wait_time += 2
                audio_file = genai.get_file(audio_file.name)
            
            if audio_file.state.name == "FAILED":
                return None, "❌ Procesarea audio a eșuat"
            
            if progress_callback:
                progress_callback(0.7, "🤖 Transcriere audio...")
            
            model = genai.GenerativeModel('gemini-1.5-pro')
            
            source = LANGUAGES.get(source_lang, "auto")
            target = LANGUAGES.get(target_lang, "Romanian")
            
            prompt = f"""
Transcrie complet acest fișier audio.

INSTRUCȚIUNI:
1. Limba sursă: {source}
2. Limba țintă: {target}
3. Transcrie TOT conținutul
4. Include timestamps aproximative [MM:SS]
5. {'TRADUCE în ' + target if source != target and source != 'auto' else 'Menține limba originală'}

Formatare:
[MM:SS] Text transcris
"""
            
            response = model.generate_content([audio_file, prompt])
            
            # Cleanup
            genai.delete_file(audio_file.name)
            
            if progress_callback:
                progress_callback(1.0, "✅ Transcriere completă!")
            
            return response.text, None
            
        else:
            # Pentru fișiere video
            if file_size_mb > GEMINI_DIRECT_UPLOAD_LIMIT_MB:
                return None, f"❌ Fișier prea mare ({file_size_mb:.1f}MB). Limita: {GEMINI_DIRECT_UPLOAD_LIMIT_MB}MB"
            
            if progress_callback:
                progress_callback(0.3, "📤 Încărcare video...")
            
            video_file = genai.upload_file(path=file_path)
            
            if progress_callback:
                progress_callback(0.5, "⏳ Procesare video...")
            
            # Așteaptă procesarea
            max_wait = 180
            wait_time = 0
            while video_file.state.name == "PROCESSING" and wait_time < max_wait:
                time.sleep(2)
                wait_time += 2
                video_file = genai.get_file(video_file.name)
                
                if progress_callback:
                    progress = 0.5 + (0.2 * (wait_time / max_wait))
                    progress_callback(progress, f"⏳ Procesare... ({wait_time}s)")
            
            if video_file.state.name == "FAILED":
                return None, "❌ Procesarea video a eșuat"
            
            if progress_callback:
                progress_callback(0.8, "🤖 Transcriere video...")
            
            model = genai.GenerativeModel('gemini-1.5-pro')
            
            source = LANGUAGES.get(source_lang, "auto")
            target = LANGUAGES.get(target_lang, "Romanian")
            
            prompt = f"""
Analizează acest video și transcrie tot conținutul audio/vocal.

INSTRUCȚIUNI:
1. Limba sursă: {source} {'(detectează automat)' if source == 'auto' else ''}
2. Limba țintă: {target}
3. Transcrie COMPLET tot dialogul
4. Include timestamps [MM:SS]
5. {'TRADUCE în ' + target if source != target and source != 'auto' else 'Menține limba originală'}
6. Notează și sunete/muzică relevante între [paranteze]

FORMAT:
[MM:SS] Text transcris
[MM:SS] [muzică de fundal]
[MM:SS] Continuare dialog...
"""
            
            generation_config = genai.types.GenerationConfig(
                temperature=0.3,
                max_output_tokens=8192,
            )
            
            response = model.generate_content(
                [video_file, prompt],
                generation_config=generation_config,
                request_options={"timeout": 600}
            )
            
            # Cleanup
            genai.delete_file(video_file.name)
            
            if progress_callback:
                progress_callback(1.0, "✅ Transcriere completă!")
            
            return response.text, None
            
    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower():
            return None, "❌ Quota API depășită"
        else:
            return None, f"❌ Eroare procesare: {error_msg}"

# ==================== WORD EXPORT ====================

def create_word_document(transcription, video_name, source_lang, target_lang, 
                        file_size_mb=0, source_type="upload", source_url=""):
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        title = doc.add_heading('Transcriere Video', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        info = doc.add_paragraph()
        info.add_run('Informații Document\n').bold = True
        info.add_run(f'📹 Video: {video_name}\n')
        
        if source_type != 'upload':
            source_icon = {'youtube': '🎬', 'gdrive': '☁️', 'direct': '🔗'}.get(source_type, '📎')
            info.add_run(f'{source_icon} Sursă: {source_type.upper()}\n')
            if source_url:
                info.add_run(f'🔗 URL: {source_url[:50]}...\n' if len(source_url) > 50 else f'🔗 URL: {source_url}\n')
        
        if file_size_mb > 0:
            info.add_run(f'📊 Dimensiune: {file_size_mb:.1f}MB\n')
        
        info.add_run(f'🌐 Limba sursă: {source_lang}\n')
        info.add_run(f'🎯 Limba țintă: {target_lang}\n')
        info.add_run(f'📅 Data: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n')
        
        doc.add_paragraph('─' * 60)
        
        doc.add_heading('Conținut Transcris', level=1)
        
        for line in transcription.split('\n'):
            if line.strip():
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()
        doc.add_paragraph('─' * 60)
        footer = doc.add_paragraph()
        footer.add_run('Generat cu AI Video Transcriber - Powered by Google Gemini').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
    except Exception as e:
        st.error(f"Eroare creare document: {e}")
        return None

# ==================== UI COMPONENTS ====================

def render_sidebar():
    with st.sidebar:
        st.markdown("## ⚙️ Configurare")
        
        st.markdown(f"""
        <div class="session-info">
            <strong>📋 ID Sesiune:</strong> {st.session_state.session_id}<br>
            <strong>🔗 Link permanent:</strong><br>
            <code>?session={st.session_state.session_id}</code><br>
            <small>💡 Salvează pentru a reveni</small>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("### 🔑 Chei API")
        
        keys = get_api_keys_from_secrets()
        
        if not keys:
            st.error("❌ Fără chei în Secrets!")
            st.markdown("**Adaugă temporar:**")
            temp_key = st.text_input("API Key:", type="password", key="temp_api")
            if st.button("Adaugă", key="add_temp"):
                if temp_key:
                    valid, msg = test_api_key(temp_key)
                    if valid:
                        if 'temp_api_keys' not in st.session_state:
                            st.session_state.temp_api_keys = []
                        st.session_state.temp_api_keys.append(temp_key)
                        st.success("✅ Adăugată!")
                        st.rerun()
                    else:
                        st.error(msg)
        else:
            st.success(f"✅ {len(keys)} chei")
        
        if 'temp_api_keys' in st.session_state:
            st.info(f"📌 {len(st.session_state.temp_api_keys)} temporare")
        
        st.markdown("---")
        
        # Capabilități
        st.markdown("""
        <div class="warning-box">
            <strong>✨ Surse Suportate:</strong><br>
            • 🎬 YouTube (max 2h)<br>
            • ☁️ Google Drive<br>
            • 🔗 Link direct video<br>
            • 📤 Upload (max 200MB)<br>
        </div>
        """, unsafe_allow_html=True)
        
        # Status biblioteci
        with st.expander("📚 Status biblioteci"):
            libs = {
                "Gemini AI": GEMINI_AVAILABLE,
                "yt-dlp": YTDLP_AVAILABLE,
                "Google API": GDRIVE_AVAILABLE,
                "python-docx": DOCX_AVAILABLE
            }
            
            for lib, status in libs.items():
                if status:
                    st.success(f"✅ {lib}")
                else:
                    st.warning(f"⚠️ {lib}")
        
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Reset", use_container_width=True):
                delete_session_data(st.session_state.session_id)
                st.session_state.messages = []
                st.session_state.transcriptions = []
                st.success("✅ Resetat!")
                st.rerun()
        
        with col2:
            if st.button("🆕 Nou", use_container_width=True):
                new_id = generate_session_id()
                create_session(new_id)
                st.session_state.session_id = new_id
                st.session_state.messages = []
                st.session_state.transcriptions = []
                st.session_state.session_loaded = True
                set_session_id_in_url(new_id)
                st.rerun()

def render_upload_tab():
    # Selector tip input
    input_type = st.radio(
        "🎯 Alege sursa video:",
        ["📤 Upload Fișier", "🔗 URL/Link", "🎬 YouTube", "☁️ Google Drive"],
        horizontal=True,
        key="input_type"
    )
    
    st.markdown("---")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        video_source = None
        video_info = {}
        
        if input_type == "📤 Upload Fișier":
            st.markdown("### 📤 Încarcă Video")
            
            uploaded_file = st.file_uploader(
                f"Selectează video (max {GEMINI_DIRECT_UPLOAD_LIMIT_MB}MB)",
                type=SUPPORTED_FORMATS,
                key="file_upload"
            )
            
            if uploaded_file:
                file_size_mb = uploaded_file.size / (1024 * 1024)
                
                if file_size_mb > GEMINI_DIRECT_UPLOAD_LIMIT_MB:
                    st.error(f"❌ Prea mare: {file_size_mb:.1f}MB (max {GEMINI_DIRECT_UPLOAD_LIMIT_MB}MB)")
                else:
                    st.success(f"✅ {uploaded_file.name} ({file_size_mb:.1f}MB)")
                    if file_size_mb <= 50:
                        st.video(uploaded_file)
                    video_source = ('upload', uploaded_file, file_size_mb)
        
        elif input_type == "🔗 URL/Link":
            st.markdown("### 🔗 URL Video Direct")
            
            video_url = st.text_input(
                "Introdu URL video direct:",
                placeholder="https://example.com/video.mp4",
                key="direct_url"
            )
            
            if video_url:
                url_type = detect_url_type(video_url)
                
                if url_type == 'youtube':
                    st.info("💡 Folosește tab-ul YouTube pentru link-uri YouTube")
                elif url_type == 'gdrive':
                    st.info("💡 Folosește tab-ul Google Drive pentru link-uri Drive")
                elif url_type == 'direct':
                    st.success(f"✅ Link video detectat")
                    video_source = ('direct', video_url, 0)
                else:
                    st.warning("⚠️ Verifică să fie un link direct către fișier video")
        
        elif input_type == "🎬 YouTube":
            st.markdown("### 🎬 YouTube Video")
            
            youtube_url = st.text_input(
                "URL sau ID video YouTube:",
                placeholder="https://www.youtube.com/watch?v=... sau ID video",
                key="youtube_url"
            )
            
            if youtube_url:
                # Extrage ID
                if 'youtube.com' in youtube_url or 'youtu.be' in youtube_url:
                    video_id = extract_video_id_youtube(youtube_url)
                else:
                    video_id = youtube_url.strip()
                
                if video_id:
                    st.success(f"✅ Video ID: {video_id}")
                    
                    # Info video
                    with st.spinner("Obțin informații..."):
                        info, error = get_youtube_info(video_id)
                        
                        if info:
                            st.markdown(f"""
                            <div class="feature-card">
                                <strong>📹 {info['title']}</strong><br>
                                👤 {info['uploader']}<br>
                                ⏱️ {info['duration']//60}:{info['duration']%60:02d}<br>
                                👁️ {info['view_count']:,} vizualizări
                            </div>
                            """, unsafe_allow_html=True)
                            
                            if info['duration'] > YOUTUBE_MAX_DURATION_MINUTES * 60:
                                st.warning(f"⚠️ Video prea lung ({info['duration']//60} min). Max: {YOUTUBE_MAX_DURATION_MINUTES} min")
                            else:
                                video_source = ('youtube', video_id, 0)
                                video_info = info
                        elif not YTDLP_AVAILABLE:
                            st.error("❌ yt-dlp nu este instalat. Adaugă în requirements.txt")
                            video_source = ('youtube', video_id, 0)  # Încearcă oricum
                        else:
                            st.error(f"❌ Eroare: {error}")
                else:
                    st.error("❌ Nu am putut extrage ID-ul video")
        
        elif input_type == "☁️ Google Drive":
            st.markdown("### ☁️ Google Drive")
            
            gdrive_url = st.text_input(
                "URL Google Drive:",
                placeholder="https://drive.google.com/file/d/.../view",
                key="gdrive_url"
            )
            
            if gdrive_url:
                file_id = extract_file_id_gdrive(gdrive_url)
                
                if file_id:
                    st.success(f"✅ File ID: {file_id[:20]}...")
                    st.info("📌 Asigură-te că fișierul este public sau 'Anyone with link'")
                    video_source = ('gdrive', file_id, 0)
                else:
                    st.error("❌ Nu am putut extrage ID-ul fișierului")
    
    with col2:
        st.markdown("### 🌐 Setări Limbă")
        
        source_lang = st.selectbox(
            "Limba sursă",
            options=list(LANGUAGES.keys()),
            index=list(LANGUAGES.keys()).index("Auto-detect"),
            key="src_lang"
        )
        
        target_lang = st.selectbox(
            "Limba țintă",
            options=[k for k in LANGUAGES.keys() if k != "Auto-detect"],
            index=0,
            key="tgt_lang"
        )
        
        if video_source:
            # Estimare timp
            source_type = video_source[0]
            est_times = {
                'upload': "1-3 minute",
                'youtube': "2-5 minute",
                'gdrive': "2-5 minute",
                'direct': "2-5 minute"
            }
            st.info(f"⏱️ Estimare: {est_times.get(source_type, '2-5 minute')}")
    
    # Buton transcriere
    if video_source:
        if st.button("🚀 Începe Transcrierea", use_container_width=True, type="primary"):
            keys = get_api_keys_from_secrets()
            
            if 'temp_api_keys' in st.session_state:
                keys = st.session_state.temp_api_keys + keys
            
            if not keys:
                st.error("❌ Nu există chei API!")
                return
            
            working_key, _, msg = get_working_api_key(keys)
            
            if not working_key:
                st.error(f"❌ {msg}")
                return
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            def update_progress(value, text):
                progress_bar.progress(min(value, 1.0))
                status_text.text(text)
            
            try:
                source_type, source_data, file_size_mb = video_source
                
                # Procesare în funcție de tip
                if source_type == 'upload':
                    # Upload direct
                    update_progress(0.1, "📁 Salvare fișier...")
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{source_data.name.split('.')[-1]}") as tmp:
                        tmp.write(source_data.getbuffer())
                        file_path = tmp.name
                    
                    video_name = source_data.name
                    source_url = ""
                    
                elif source_type == 'youtube':
                    # Descarcă de pe YouTube
                    file_path, video_name, download_type = download_youtube_video(
                        source_data, 
                        max_size_mb=GEMINI_DIRECT_UPLOAD_LIMIT_MB,
                        progress_callback=update_progress
                    )
                    
                    if not file_path:
                        st.error(f"❌ Nu s-a putut descărca: {download_type}")
                        return
                    
                    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                    source_url = f"https://youtube.com/watch?v={source_data}"
                    
                    # Flag pentru audio-only
                    is_audio_only = (download_type == 'audio_only')
                
                elif source_type == 'gdrive':
                    # Descarcă de pe Google Drive
                    file_path, video_name, error = download_gdrive_video(
                        source_data,
                        progress_callback=update_progress
                    )
                    
                    if not file_path:
                        st.error(f"❌ Nu s-a putut descărca: {error}")
                        return
                    
                    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                    source_url = f"https://drive.google.com/file/d/{source_data}"
                
                elif source_type == 'direct':
                    # Descarcă de la URL direct
                    file_path, video_name, error = download_direct_video(
                        source_data,
                        progress_callback=update_progress
                    )
                    
                    if not file_path:
                        st.error(f"❌ Nu s-a putut descărca: {error}")
                        return
                    
                    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                    source_url = source_data
                
                # Verifică dimensiunea
                if file_size_mb > GEMINI_DIRECT_UPLOAD_LIMIT_MB:
                    st.error(f"❌ Fișier prea mare după descărcare: {file_size_mb:.1f}MB")
                    os.unlink(file_path)
                    return
                
                # Procesare și transcriere
                is_audio = source_type == 'youtube' and 'is_audio_only' in locals() and is_audio_only
                
                transcription, error = process_and_transcribe(
                    file_path, 
                    source_lang, 
                    target_lang,
                    working_key,
                    file_size_mb,
                    update_progress,
                    is_audio_only=is_audio
                )
                
                # Cleanup
                try:
                    os.unlink(file_path)
                except:
                    pass
                
                if error:
                    st.error(error)
                    return
                
                if not transcription:
                    st.error("❌ Nu s-a putut genera transcrierea")
                    return
                
                # Salvează în DB
                save_transcription(
                    st.session_state.session_id,
                    video_name,
                    source_lang,
                    target_lang,
                    transcription,
                    file_size_mb,
                    source_type,
                    source_url,
                    source_type
                )
                
                update_progress(1.0, "✅ Transcriere completă!")
                st.success(f"🎉 Video transcris cu succes!")
                
                # Afișează rezultatul
                st.markdown("### 📝 Transcriere")
                st.markdown(f"""
                <div class="transcription-box">
{transcription}
                </div>
                """, unsafe_allow_html=True)
                
                # Butoane descărcare
                col1, col2 = st.columns(2)
                
                with col1:
                    word_doc = create_word_document(
                        transcription, 
                        video_name,
                        source_lang, 
                        target_lang,
                        file_size_mb,
                        source_type,
                        source_url
                    )
                    if word_doc:
                        st.download_button(
                            "📥 Descarcă Word",
                            word_doc,
                            f"transcriere_{video_name.split('.')[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                
                with col2:
                    st.download_button(
                        "📥 Descarcă Text",
                        transcription,
                        f"transcriere_{video_name.split('.')[0]}.txt",
                        mime="text/plain"
                    )
                
            except Exception as e:
                st.error(f"❌ Eroare: {str(e)}")

def render_history_tab():
    st.markdown("### 📜 Istoric Transcrieri")
    
    transcriptions = get_transcriptions(st.session_state.session_id)
    
    if not transcriptions:
        st.info("📭 Nu există transcrieri încă")
    else:
        for i, trans in enumerate(transcriptions):
            # Icon pentru tip sursă
            source_icon = {
                'upload': '📤',
                'youtube': '🎬',
                'gdrive': '☁️',
                'direct': '🔗'
            }.get(trans.get('source_type', 'upload'), '📎')
            
            title = f"{source_icon} {trans['video_name']}"
            if trans.get('file_size_mb', 0) > 0:
                title += f" ({trans['file_size_mb']:.1f}MB)"
            title += f" - {trans['created_at']}"
            
            with st.expander(title, expanded=False):
                # Info
                cols = st.columns(4)
                with cols[0]:
                    st.write(f"**Sursă:** {trans.get('source_language', 'N/A')}")
                with cols[1]:
                    st.write(f"**Țintă:** {trans.get('target_language', 'N/A')}")
                with cols[2]:
                    st.write(f"**Tip:** {trans.get('source_type', 'upload')}")
                with cols[3]:
                    st.write(f"**Status:** {trans.get('status', 'completed')}")
                
                # URL sursă dacă există
                if trans.get('source_url'):
                    st.caption(f"🔗 {trans['source_url'][:100]}...")
                
                # Transcriere
                st.text_area(
                    "Transcriere",
                    trans['transcription'],
                    height=300,
                    key=f"hist_{trans['id']}",
                    label_visibility="collapsed"
                )
                
                # Descărcări
                col1, col2 = st.columns(2)
                with col1:
                    word_doc = create_word_document(
                        trans['transcription'],
                        trans['video_name'],
                        trans.get('source_language', ''),
                        trans.get('target_language', ''),
                        trans.get('file_size_mb', 0),
                        trans.get('source_type', ''),
                        trans.get('source_url', '')
                    )
                    if word_doc:
                        st.download_button(
                            "📥 Word",
                            word_doc,
                            f"trans_{trans['id']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"w_{trans['id']}"
                        )
                
                with col2:
                    st.download_button(
                        "📥 Text",
                        trans['transcription'],
                        f"trans_{trans['id']}.txt",
                        mime="text/plain",
                        key=f"t_{trans['id']}"
                    )

def render_chat_tab():
    st.markdown("### 💬 Chat cu AI despre Transcrieri")
    
    messages = get_messages(st.session_state.session_id)
    
    for msg in messages:
        with st.chat_message(msg['role']):
            st.markdown(msg['content'])
    
    if prompt := st.chat_input("Întreabă despre transcrieri..."):
        with st.chat_message("user"):
            st.markdown(prompt)
        save_message(st.session_state.session_id, "user", prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("Generez răspuns..."):
                keys = get_api_keys_from_secrets()
                if 'temp_api_keys' in st.session_state:
                    keys = st.session_state.temp_api_keys + keys
                
                if not keys:
                    st.error("❌ Nu există chei API!")
                    return
                
                working_key, _, _ = get_working_api_key(keys)
                if not working_key:
                    st.error("❌ Toate cheile sunt invalide!")
                    return
                
                try:
                    genai.configure(api_key=working_key)
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    
                    # Context
                    recent = get_transcriptions(st.session_state.session_id)[:2]
                    context = ""
                    if recent:
                        context = "Transcrieri recente:\n"
                        for t in recent:
                            source_type = t.get('source_type', 'upload')
                            context += f"- {t['video_name']} ({source_type}): {t['transcription'][:300]}...\n\n"
                    
                    full_prompt = f"""
{context}

Utilizator: {prompt}

Răspunde în română. Dacă întrebarea e despre transcrieri, folosește contextul de mai sus.
"""
                    
                    response = model.generate_content(full_prompt)
                    response_text = response.text
                    
                    st.markdown(response_text)
                    save_message(st.session_state.session_id, "assistant", response_text)
                    
                except Exception as e:
                    st.error(f"❌ Eroare: {e}")

# ==================== MAIN ====================

def main():
    init_database()
    init_session()
    
    render_sidebar()
    
    st.markdown("""
    <div class="main-header">
        <h1>🎬 AI Video Transcriber</h1>
        <p>Transcrie videouri din YouTube, Google Drive, link-uri directe sau upload</p>
        <p style="font-size: 0.9rem; opacity: 0.8;">
            Powered by Google Gemini AI | Suport multilingv | Export Word/Text
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    if not GEMINI_AVAILABLE:
        st.error("❌ Google Generative AI nu este instalat!")
        st.stop()
    
    # Tabs principale
    tab1, tab2, tab3 = st.tabs(["🚀 Transcriere", "📜 Istoric", "💬 Chat AI"])
    
    with tab1:
        render_upload_tab()
    
    with tab2:
        render_history_tab()
    
    with tab3:
        render_chat_tab()

if __name__ == "__main__":
    main()
